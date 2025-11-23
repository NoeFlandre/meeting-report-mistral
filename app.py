import streamlit as st
import io
from mistralai import Mistral
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydub import AudioSegment
import os
from datetime import date


# ==================== CONFIGURATION ====================
st.set_page_config(
    page_title="Générateur de Compte Rendu",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personnalisé avec meilleurs contrastes
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 0;
    }
    .subtitle {
        text-align: center;
        color: #555;
        margin-bottom: 2rem;
        font-size: 1.1rem;
    }
    .step-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1.2rem;
        border-radius: 0.5rem;
        margin: 1.5rem 0;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .step-header h3 {
        color: white !important;
        margin: 0;
        font-weight: 600;
    }
    .success-box {
        background-color: #d4edda;
        border-left: 5px solid #28a745;
        padding: 1rem;
        border-radius: 0.3rem;
        margin: 1rem 0;
        color: #155724;
    }
    .warning-box {
        background-color: #fff3cd;
        border-left: 5px solid #ffc107;
        padding: 1rem;
        border-radius: 0.3rem;
        margin: 1rem 0;
        color: #856404;
    }
    </style>
""", unsafe_allow_html=True)


# ==================== FONCTIONS UTILITAIRES ====================

def get_mistral_client(key):
    """Initialise le client Mistral avec validation"""
    if not key or len(key) < 10:
        st.error("⚠️ Veuillez entrer une clé API Mistral valide dans la barre latérale.")
        st.stop()
    return Mistral(api_key=key)


def chunk_audio(audio_file, chunk_duration_minutes=10):
    """
    Divise le fichier audio en chunks pour éviter les limites de contexte.
    Retourne: (liste de chunks, extension du fichier)
    """
    audio = AudioSegment.from_file(io.BytesIO(audio_file.getvalue()))
    file_ext = os.path.splitext(audio_file.name)[1].lower().replace('.', '')
    
    chunk_duration_ms = chunk_duration_minutes * 60 * 1000
    total_duration_ms = len(audio)
    
    # Si l'audio est court, retourner directement
    if total_duration_ms <= chunk_duration_ms:
        return [(audio, 0, total_duration_ms / 1000 / 60)], file_ext
    
    # Sinon, découper en chunks
    chunks = []
    for i in range(0, total_duration_ms, chunk_duration_ms):
        chunk = audio[i:i + chunk_duration_ms]
        start_time = i / 1000 / 60
        end_time = min((i + chunk_duration_ms) / 1000 / 60, total_duration_ms / 1000 / 60)
        chunks.append((chunk, start_time, end_time))
    
    return chunks, file_ext


def transcribe_chunk(client, chunk_audio, file_name, file_ext):
    """Transcrit un chunk audio unique"""
    buffer = io.BytesIO()
    chunk_audio.export(buffer, format=file_ext)
    chunk_bytes = buffer.getvalue()
    
    response = client.audio.transcriptions.complete(
        model="voxtral-mini-latest",
        file={
            "file_name": file_name,
            "content": chunk_bytes,
        }
    )
    
    return response.text


def generate_meeting_notes(client, transcript, municipalite, sujets):
    """Génère un compte rendu structuré avec Mistral Medium"""
    
    context_sujets = f"\n\nSujets prévus à l'ordre du jour: {sujets}" if sujets else ""
    
    system_prompt = """Tu es un expert en rédaction de comptes rendus de réunion pour collectivités territoriales françaises.

À partir de la transcription fournie, génère un compte rendu professionnel et structuré avec:

## 1. INFORMATIONS GÉNÉRALES
- Date et heure (si mentionnées)
- Participants présents

## 2. ORDRE DU JOUR
Liste des points discutés

## 3. DISCUSSIONS ET DÉBATS
Pour chaque point à l'ordre du jour:
- Résumé des discussions
- Positions exprimées
- Arguments principaux

## 4. DÉCISIONS PRISES
Liste claire des décisions votées ou approuvées avec:
- La décision
- Vote (si mentionné: pour/contre/abstention)

## 5. ACTIONS À RÉALISER
Pour chaque action:
- Description de l'action
- Responsable (si mentionné)
- Échéance (si mentionnée)

## 6. PROCHAINE RÉUNION
Date et sujets prévus (si mentionnés)

Consignes:
- Utilise un français professionnel et clair
- Sois concis mais complet
- Utilise des listes à puces pour la lisibilité
- Si une information n'est pas dans la transcription, note "Non précisé"
- Format: Markdown avec titres clairs (##, ###)"""

    response = client.chat.complete(
        model="mistral-medium-latest",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Municipalité: {municipalite}{context_sujets}\n\nTranscription de la réunion:\n\n{transcript}"}
        ],
        max_tokens=4000
    )
    
    return response.choices[0].message.content


def create_docx_report(meeting_notes, municipalite, date_reunion, sujets):
    """Crée un document DOCX professionnel"""
    document = Document()
    
    # En-tête
    heading = document.add_heading('COMPTE RENDU DE RÉUNION', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Info municipalité
    info_para = document.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f"{municipalite}\n{date_reunion.strftime('%d/%m/%Y')}")
    info_run.font.size = Pt(12)
    info_run.font.italic = True
    
    document.add_paragraph()  # Espacement
    
    # Ligne de séparation
    document.add_paragraph('_' * 80)
    
    # Sujets
    if sujets:
        document.add_heading("Sujets traités :", level=2)
        document.add_paragraph(sujets)
        document.add_paragraph()
    
    # Contenu du compte rendu
    lines = meeting_notes.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            document.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            text = line.split('. ', 1)[1] if '. ' in line else line
            document.add_paragraph(text, style='List Number')
        else:
            if line and not line.startswith('_'):
                document.add_paragraph(line)
    
    # Pied de page
    document.add_paragraph()
    document.add_paragraph('_' * 80)
    footer = document.add_paragraph("Document généré automatiquement - Compte Rendu IA")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    return document


# ==================== INTERFACE PRINCIPALE ====================

# En-tête
st.markdown('<p class="main-header">📋 Générateur de Compte Rendu de Réunion</p>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">Solution automatique pour secrétaires de mairie</p>', unsafe_allow_html=True)

# Barre latérale - Configuration
with st.sidebar:
    st.header("⚙️ Configuration")
    
    with st.expander("🔑 Clé API Mistral", expanded=False):
        # Auto-use secret from Streamlit Cloud
        if "MISTRAL_API_KEY" in st.secrets:
            api_key = st.secrets["MISTRAL_API_KEY"]
            st.success("✅ Clé API configurée automatiquement")
        else:
            # Fallback for local development
            api_key = st.text_input(
                "Entrez votre clé API",
                type="password",
                help="Obtenez votre clé sur console.mistral.ai"
            )
            if api_key:
                st.success("✅ Clé API configurée manuellement")
    
    st.divider()
    
    with st.expander("⚡ Paramètres avancés"):
        chunk_duration = st.slider(
            "Durée max par segment (minutes)",
            min_value=5,
            max_value=15,
            value=10,
            help="Les fichiers longs sont découpés automatiquement"
        )
        
    st.divider()
    
    # Aide
    with st.expander("❓ Aide"):
        st.markdown("""
        **Comment utiliser cette application :**
        
        1. Entrez votre clé API Mistral
        2. Remplissez les informations de la réunion
        3. Téléchargez votre fichier audio
        4. Cliquez sur "Générer le compte rendu"
        5. Téléchargez le document DOCX
        
        **Formats audio acceptés :**
        - MP3
        - WAV
        - M4A
        
        **Durée max recommandée :** 2 heures
        """)

# ==================== FORMULAIRE PRINCIPAL ====================

st.markdown('<div class="step-header"><h3>📝 Étape 1 : Informations de la réunion</h3></div>', unsafe_allow_html=True)

with st.form("meeting_info_form", clear_on_submit=False):
    col1, col2 = st.columns(2)
    
    with col1:
        municipalite = st.text_input(
            "Nom de la municipalité *",
            value="Mairie de",
            help="Ex: Mairie de Paris 15ème"
        )
        
    with col2:
        date_reunion = st.date_input(
            "Date de la réunion *",
            value=date.today(),
            help="Sélectionnez la date de la réunion"
        )
    
    sujets = st.text_area(
        "Ordre du jour (optionnel)",
        placeholder="Ex: Budget 2025, Travaux voirie, Festivités estivales...",
        height=100,
        help="Listez les sujets prévus à l'ordre du jour"
    )
    
    form_submitted = st.form_submit_button(
        "✅ Valider les informations",
        use_container_width=True,
        type="primary"
    )
    
    if form_submitted:
        if municipalite and date_reunion:
            st.session_state['meeting_info'] = {
                'municipalite': municipalite,
                'date_reunion': date_reunion,
                'sujets': sujets
            }
            st.success("✅ Informations enregistrées !")
        else:
            st.error("⚠️ Veuillez remplir tous les champs obligatoires (*)")

# ==================== UPLOAD AUDIO ====================

st.markdown('<div class="step-header"><h3>🎙️ Étape 2 : Fichier audio de la réunion</h3></div>', unsafe_allow_html=True)

audio_file = st.file_uploader(
    "Téléchargez l'enregistrement audio",
    type=['mp3', 'wav', 'm4a'],
    help="Formats acceptés: MP3, WAV, M4A"
)

if audio_file:
    # Afficher les infos du fichier
    col1, col2, col3 = st.columns(3)
    
    try:
        audio_temp = AudioSegment.from_file(io.BytesIO(audio_file.getvalue()))
        duration_minutes = len(audio_temp) / 1000 / 60
        file_size_mb = len(audio_file.getvalue()) / (1024 * 1024)
        
        with col1:
            st.metric("📁 Fichier", audio_file.name)
        with col2:
            st.metric("⏱️ Durée", f"{duration_minutes:.1f} min")
        with col3:
            st.metric("💾 Taille", f"{file_size_mb:.1f} MB")
        
        # Lecteur audio
        st.audio(audio_file)
        
    except Exception as e:
        st.error(f"❌ Erreur lors de la lecture du fichier audio: {str(e)}")

# ==================== GÉNÉRATION ====================

st.markdown('<div class="step-header"><h3>🚀 Étape 3 : Génération du compte rendu</h3></div>', unsafe_allow_html=True)

# Vérifications avant génération
can_generate = all([
    audio_file is not None,
    api_key,
    'meeting_info' in st.session_state
])

if not can_generate:
    missing = []
    if not api_key:
        missing.append("Clé API Mistral")
    if 'meeting_info' not in st.session_state:
        missing.append("Informations de réunion")
    if not audio_file:
        missing.append("Fichier audio")
    
    st.warning(f"⚠️ Éléments manquants : {', '.join(missing)}")

if st.button(
    "🚀 Générer le compte rendu",
    use_container_width=True,
    type="primary",
    disabled=not can_generate
):
    
    client = get_mistral_client(api_key)
    meeting_info = st.session_state['meeting_info']
    
    try:
        # ========== TRANSCRIPTION ==========
        st.markdown("### 📝 Transcription en cours...")
        
        with st.spinner("Analyse du fichier audio..."):
            chunks, file_ext = chunk_audio(audio_file, chunk_duration)
        
        st.info(f"🔪 Fichier divisé en {len(chunks)} segment(s)")
        
        # Barre de progression
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        full_transcript = []
        
        for idx, (chunk, start_time, end_time) in enumerate(chunks):
            status_text.text(f"Transcription du segment {idx+1}/{len(chunks)} ({start_time:.1f}-{end_time:.1f} min)...")
            
            chunk_text = transcribe_chunk(
                client,
                chunk,
                f"chunk_{idx}.{file_ext}",
                file_ext
            )
            
            full_transcript.append(chunk_text)
            progress_bar.progress((idx + 1) / len(chunks))
        
        transcript = "\n\n".join(full_transcript)
        status_text.empty()
        st.success("✅ Transcription terminée !")
        
        # ========== GÉNÉRATION COMPTE RENDU ==========
        st.markdown("### 🤖 Génération du compte rendu...")
        
        with st.spinner("Analyse et rédaction par Mistral Medium..."):
            meeting_notes = generate_meeting_notes(
                client,
                transcript,
                meeting_info['municipalite'],
                meeting_info['sujets']
            )
        
        st.success("✅ Compte rendu généré !")
        
        # ========== APERÇU ==========
        st.markdown("### 👀 Aperçu du compte rendu")
        
        with st.expander("📄 Voir le compte rendu", expanded=True):
            st.markdown(meeting_notes)
        
        # ========== CRÉATION DOCX ==========
        with st.spinner("Création du document Word..."):
            doc = create_docx_report(
                meeting_notes,
                meeting_info['municipalite'],
                meeting_info['date_reunion'],
                meeting_info['sujets']
            )
            
            # Conversion en bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
        
        # ========== TÉLÉCHARGEMENT ==========
        st.markdown('<div class="success-box">', unsafe_allow_html=True)
        st.markdown("### ✅ Document prêt !")
        
        filename = f"CR_{meeting_info['municipalite'].replace(' ', '_')}_{meeting_info['date_reunion'].strftime('%Y%m%d')}.docx"
        
        col1, col2 = st.columns([3, 1])
        
        with col1:
            st.download_button(
                label="📥 Télécharger le compte rendu (DOCX)",
                data=doc_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col2:
            # Bouton pour télécharger aussi la transcription
            st.download_button(
                label="📄 Transcription",
                data=transcript,
                file_name=f"Transcription_{meeting_info['date_reunion'].strftime('%Y%m%d')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # ========== STATS ==========
        with st.expander("📊 Statistiques"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Mots transcrits", len(transcript.split()))
            with col2:
                st.metric("Segments traités", len(chunks))
            with col3:
                st.metric("Durée totale", f"{duration_minutes:.1f} min")
        
    except Exception as e:
        st.error(f"❌ Erreur lors de la génération : {str(e)}")
        st.info("💡 Conseil : Vérifiez votre clé API et la qualité du fichier audio")

# ==================== PIED DE PAGE ==========
st.divider()
st.markdown("""
<div style='text-align: center; color: #666; font-size: 0.9rem;'>
    <p>🔒 Vos données sont traitées de manière sécurisée et ne sont pas conservées</p>
    <p>Développé avec Streamlit • Propulsé par Mistral AI</p>
</div>
""", unsafe_allow_html=True)
